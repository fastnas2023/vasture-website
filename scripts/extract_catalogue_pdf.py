#!/usr/bin/env python3
"""Extract catalogue PDF text, OCR, embedded images, previews and editable DOCX."""

from __future__ import annotations

import argparse
from concurrent.futures import ThreadPoolExecutor
import hashlib
import json
import re
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


BLUE = RGBColor(30, 79, 111)
DARK_BLUE = RGBColor(16, 42, 61)
MUTED = RGBColor(90, 103, 112)


def run(command: list[str]) -> None:
    subprocess.run(command, check=True)


def clean_ocr(text: str) -> str:
    lines = [line.rstrip() for line in text.replace("\x0c", "").splitlines()]
    cleaned: list[str] = []
    blank = False
    for line in lines:
        if line.strip():
            cleaned.append(line.strip())
            blank = False
        elif cleaned and not blank:
            cleaned.append("")
            blank = True
    return "\n".join(cleaned).strip()


def ocr_page(preview: Path) -> str:
    result = subprocess.run(
        ["tesseract", str(preview), "stdout", "-l", "eng", "--psm", "11"],
        check=True,
        capture_output=True,
        text=True,
    )
    return clean_ocr(result.stdout)


def safe_name(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip("-.")
    return value or "image.bin"


def set_font(run, name: str = "Calibri", size: float | None = None,
             color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_font(run, size=9, color=MUTED)


def build_docx(output: Path, source_name: str, page_records: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.text = "MODISH / CENTRO EDILE - Editable OCR Extract"
    set_font(header.runs[0], size=9, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(72)
    title.paragraph_format.space_after = Pt(10)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("MODISH / CENTRO EDILE"), size=26, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(subtitle.add_run("Editable Catalogue Text Extract"), size=16, color=BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    set_font(
        meta.add_run(f"Source: {source_name}\nPages: {len(page_records)}\nExtracted: {date.today().isoformat()}"),
        size=10.5,
        color=MUTED,
    )

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(16)
    note.paragraph_format.space_after = Pt(10)
    set_font(note.add_run("Editing note\n"), size=11, color=DARK_BLUE, bold=True)
    set_font(
        note.add_run(
            "The original catalogue text is largely outlined or embedded in page artwork. "
            "The text below was recovered with OCR and remains editable, but product names, "
            "measurements and small labels should be checked against the page previews."
        ),
        size=10.5,
        color=MUTED,
    )

    doc.add_page_break()

    for index, record in enumerate(page_records):
        doc.add_heading(f"PDF Page {record['page']:02d}", level=1)
        source_line = doc.add_paragraph()
        source_line.paragraph_format.space_after = Pt(8)
        set_font(
            source_line.add_run(
                f"Preview: {record['preview_file']} | Extracted images: {len(record['image_files'])}"
            ),
            size=9,
            color=MUTED,
        )

        doc.add_heading("OCR editable text", level=2)
        ocr_text = record["ocr_text"] or "[No OCR text detected on this page]"
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.keep_together = False
        ocr_lines = [line for line in ocr_text.splitlines() if line.strip()]
        for line_index, line in enumerate(ocr_lines):
            run = paragraph.add_run(line.strip())
            set_font(run, size=8.5)
            if line_index != len(ocr_lines) - 1:
                run.add_break()

        native_text = record["native_text"].strip()
        if native_text and len(native_text) > len(ocr_text) * 1.25:
            doc.add_heading("Native PDF text layer", level=2)
            paragraph = doc.add_paragraph()
            set_font(paragraph.add_run(native_text), size=9.5, color=MUTED)

        if index != len(page_records) - 1:
            doc.add_page_break()

    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_pdf", type=Path)
    parser.add_argument("output_dir", type=Path)
    parser.add_argument("--dpi", type=int, default=220)
    parser.add_argument("--docx", action="store_true", help="Also create an editable OCR DOCX")
    parser.add_argument(
        "--workers",
        type=int,
        default=1,
        help="Number of parallel Tesseract OCR workers (default: 1; raise only after benchmarking)",
    )
    args = parser.parse_args()

    source = args.input_pdf.resolve()
    output = args.output_dir.resolve()
    pages_dir = output / "pages"
    text_dir = output / "text" / "pages"
    images_dir = output / "images" / "extracted"
    for directory in (pages_dir, text_dir, images_dir):
        directory.mkdir(parents=True, exist_ok=True)

    page_prefix = pages_dir / "page"
    run([
        "pdftoppm", "-jpeg", "-r", str(args.dpi), "-jpegopt", "quality=88",
        str(source), str(page_prefix),
    ])

    preview_files = sorted(pages_dir.glob("page-*.jpg"))
    reader = PdfReader(str(source))
    if len(preview_files) != len(reader.pages):
        raise RuntimeError(
            f"Rendered page count {len(preview_files)} does not match PDF page count {len(reader.pages)}"
        )

    records: list[dict] = []
    image_manifest: list[dict] = []

    with ThreadPoolExecutor(max_workers=max(1, args.workers)) as executor:
        ocr_texts = list(executor.map(ocr_page, preview_files))

    for page_number, (pdf_page, preview, ocr_text) in enumerate(
        zip(reader.pages, preview_files, ocr_texts), start=1
    ):
            native_text = clean_ocr(pdf_page.extract_text() or "")

            text_file = text_dir / f"page-{page_number:03d}.txt"
            text_file.write_text(ocr_text + "\n", encoding="utf-8")

            page_images: list[str] = []
            for image_number, image_object in enumerate(list(pdf_page.images), start=1):
                original_name = safe_name(image_object.name)
                suffix = Path(original_name).suffix.lower()
                if suffix not in {".jpg", ".jpeg", ".png", ".jp2", ".jpx", ".tif", ".tiff"}:
                    suffix = ".bin"
                filename = f"page-{page_number:03d}-image-{image_number:02d}{suffix}"
                target = images_dir / filename
                target.write_bytes(image_object.data)
                digest = hashlib.sha256(image_object.data).hexdigest()
                rel = target.relative_to(output).as_posix()
                page_images.append(rel)
                image_manifest.append({
                    "page": page_number,
                    "index": image_number,
                    "file": rel,
                    "original_name": image_object.name,
                    "bytes": len(image_object.data),
                    "sha256": digest,
                })

            records.append({
                "page": page_number,
                "preview_file": preview.relative_to(output).as_posix(),
                "text_file": text_file.relative_to(output).as_posix(),
                "image_files": page_images,
                "native_text": native_text,
                "ocr_text": ocr_text,
            })

    data = {
        "source_pdf": source.name,
        "source_path": source.relative_to(source.parents[1]).as_posix(),
        "page_count": len(records),
        "extraction_date": date.today().isoformat(),
        "ocr_language": "eng",
        "ocr_engine": "tesseract --psm 11",
        "pages": records,
        "images": image_manifest,
    }
    (output / "catalogue.json").write_text(
        json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )

    markdown = [
        "# MODISH / CENTRO EDILE - Editable OCR Extract",
        "",
        f"Source: `{source.name}`",
        f"Pages: {len(records)}",
        "",
        "> OCR output is editable but must be checked against page previews for product names, measurements and small labels.",
        "",
    ]
    for record in records:
        markdown.extend([
            f"## PDF Page {record['page']:02d}",
            "",
            f"Preview: `{record['preview_file']}`",
            "",
            record["ocr_text"] or "[No OCR text detected]",
            "",
        ])
    (output / "catalogue.md").write_text("\n".join(markdown), encoding="utf-8")

    readme = f"""# PDF extraction output

Source: `{source.name}`

## Contents

- `catalogue.md`: editable plain-text catalogue in Markdown.
- `catalogue.json`: structured page text and image mappings for future backend import.
- `text/pages/`: one UTF-8 OCR text file per PDF page.
- `images/extracted/`: embedded raster images extracted from the PDF.
- `pages/`: rendered page previews used for visual comparison and OCR.

## Limitation

The PDF was exported from Adobe InDesign, but the original InDesign source is not embedded. Most product labels are outlines or page artwork, so exact editable layout cannot be reconstructed automatically. OCR text should be proofread against `pages/` before publication.
"""
    (output / "README.md").write_text(readme, encoding="utf-8")

    if args.docx:
        build_docx(output / "modish-centro-edile-en-v2-editable.docx", source.name, records)
    print(json.dumps({
        "pages": len(records),
        "images": len(image_manifest),
        "output": str(output),
    }, ensure_ascii=False))


if __name__ == "__main__":
    main()
